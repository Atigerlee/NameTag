import os
import sys
import tkinter as tk
from tkinter import messagebox, filedialog, IntVar, StringVar, Radiobutton
from docx import Document
from docx.shared import Pt, RGBColor

try:
    from docx.oxml.ns import qn
except ImportError:
    def qn(tag):
        return "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + tag.split(':')[-1]


def resource_path(relative_path):
    """取得資源的正確路徑：開發階段或 PyInstaller 打包後"""
    try:
        base_path = sys._MEIPASS  # PyInstaller 打包後會設置
    except Exception:
        base_path = os.path.abspath(".")  # 開發階段
    return os.path.join(base_path, relative_path)


def add_spaces_between_chars(text):
    return " ".join(text) if text else text


def replace_text_in_paragraph(paragraph, replacements):
    runs_format = []
    for run in paragraph.runs:
        format_data = {
            'bold': run.bold,
            'italic': run.italic,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'color': run.font.color.rgb if run.font.color else None,
            'underline': run.underline,
            'text': run.text
        }
        runs_format.append(format_data)

    full_text = ''.join(run.text for run in paragraph.runs)
    original_text = full_text

    for key, value in replacements.items():
        full_text = full_text.replace(key, value)

    if original_text == full_text:
        return

    for _ in range(len(paragraph.runs)):
        paragraph._p.remove(paragraph.runs[0]._r)

    if runs_format:
        new_run = paragraph.add_run(full_text)
        first_format = runs_format[0]
        new_run.bold = first_format['bold']
        new_run.italic = first_format['italic']
        new_run.underline = first_format['underline']

        if first_format['font_name']:
            new_run.font.name = first_format['font_name']
            try:
                if "微軟正黑體" in first_format['font_name'] or "Microsoft JhengHei" in first_format['font_name']:
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), "Microsoft JhengHei")
            except NameError:
                pass

        if first_format['font_size']:
            new_run.font.size = first_format['font_size']
        if first_format['color']:
            new_run.font.color.rgb = first_format['color']


def open_file(file_path):
    try:
        os.startfile(file_path)  # Windows
    except AttributeError:
        import subprocess
        try:
            subprocess.call(['open', file_path])  # macOS
        except:
            subprocess.call(['xdg-open', file_path])  # Linux


def generate_doc(name1, company1, name2=None, company2=None, add_spaces=False, auto_open=True, mode="two_people"):
    template_path = resource_path("template.docx")
    output_filename = f"{name1}_名牌.docx" if mode == "one_person" else f"{name1}_{name2}_名牌.docx"
    output_path = os.path.join(os.getcwd(), output_filename)

    try:
        doc = Document(template_path)
    except Exception as e:
        messagebox.showerror("錯誤", f"無法開啟範本檔案: {str(e)}")
        return False

    if add_spaces:
        name1 = add_spaces_between_chars(name1)
        company1 = add_spaces_between_chars(company1)
        if mode == "two_people":
            name2 = add_spaces_between_chars(name2)
            company2 = add_spaces_between_chars(company2)

    replacements = {
        "{{NAME1}}": name1,
        "{{COMPANY1}}": company1,
        "{{NAME2}}": name2 or "",
        "{{COMPANY2}}": company2 or ""
    }

    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, replacements)

    try:
        doc.save(output_path)
        if auto_open:
            open_file(output_path)
        messagebox.showinfo("成功", f"名牌已儲存於：\n{output_path}")
        return output_path
    except Exception as e:
        messagebox.showerror("錯誤", f"儲存檔案失敗: {str(e)}")
        return False


def run_gui():
    root = tk.Tk()
    root.title("名牌產生器")
    root.geometry("700x600")

    frame = tk.Frame(root, padx=10, pady=10)
    frame.pack(fill=tk.BOTH, expand=True)

    mode_var = StringVar(value="two_people")
    person1_name = StringVar()
    person1_company = StringVar()
    person2_name = StringVar()
    person2_company = StringVar()

    mode_frame = tk.LabelFrame(frame, text="名牌模式")
    mode_frame.pack(fill="x", pady=5)
    Radiobutton(mode_frame, text="雙人名牌", variable=mode_var, value="two_people").pack(side="left", padx=20)
    Radiobutton(mode_frame, text="單人名牌", variable=mode_var, value="one_person").pack(side="left", padx=20)

    def create_input_section(label_text, text_var1, text_var2):
        section = tk.LabelFrame(frame, text=label_text)
        section.pack(fill="x", pady=5)
        tk.Label(section, text="名字：").grid(row=0, column=0, sticky="w")
        tk.Entry(section, textvariable=text_var1, width=30).grid(row=0, column=1, padx=5, pady=5)
        tk.Label(section, text="公司：").grid(row=1, column=0, sticky="w")
        tk.Entry(section, textvariable=text_var2, width=30).grid(row=1, column=1, padx=5, pady=5)
        return section

    p1_section = create_input_section("第一個人資料", person1_name, person1_company)
    p2_section = create_input_section("第二個人資料", person2_name, person2_company)

    options_frame = tk.LabelFrame(frame, text="選項")
    options_frame.pack(fill="x", pady=5)
    add_spaces_var = IntVar()
    tk.Checkbutton(options_frame, text="字元間添加空格", variable=add_spaces_var).pack(anchor="w")
    auto_open_var = IntVar(value=1)
    tk.Checkbutton(options_frame, text="產生後自動開啟 Word 檔", variable=auto_open_var).pack(anchor="w")

    preview_label = tk.Label(frame, text="名字 預覽：")
    preview_label.pack(pady=10)

    def update_preview(*args):
        n1 = person1_name.get().strip()
        n2 = person2_name.get().strip()
        spaced = add_spaces_between_chars
        if mode_var.get() == "two_people":
            preview_label.config(text=f"名字1：{spaced(n1)}\n名字2：{spaced(n2)}" if add_spaces_var.get() else f"名字1：{n1}\n名字2：{n2}")
        else:
            preview_label.config(text=f"名字：{spaced(n1)}" if add_spaces_var.get() else f"名字：{n1}")

    for var in [person1_name, person2_name]:
        var.trace_add("write", update_preview)
    add_spaces_var.trace_add("write", update_preview)
    mode_var.trace_add("write", lambda *args: (p2_section.pack_forget() if mode_var.get() == "one_person" else p2_section.pack(fill="x", pady=5), update_preview()))

    def on_submit():
        if not person1_name.get().strip() or not person1_company.get().strip():
            messagebox.showwarning("提示", "請填寫第一個人的完整資料")
            return
        if mode_var.get() == "two_people" and (not person2_name.get().strip() or not person2_company.get().strip()):
            messagebox.showwarning("提示", "請填寫第二個人的完整資料")
            return

        generate_doc(
            person1_name.get().strip(),
            person1_company.get().strip(),
            person2_name.get().strip(),
            person2_company.get().strip(),
            add_spaces_var.get(),
            auto_open_var.get(),
            mode_var.get()
        )

    tk.Button(frame, text="產生名牌", command=on_submit).pack(pady=10)
    tk.Button(frame, text="退出", command=root.quit).pack()

    # 檢查範本是否存在
    template_path = resource_path("template.docx")
    status_label = tk.Label(frame, text=f"範本檔案: template.docx ({'存在' if os.path.exists(template_path) else '不存在'})", fg="green" if os.path.exists(template_path) else "red")
    status_label.pack(pady=5)

    update_preview()
    if mode_var.get() == "one_person":
        p2_section.pack_forget()

    root.mainloop()


if __name__ == "__main__":
    run_gui()
